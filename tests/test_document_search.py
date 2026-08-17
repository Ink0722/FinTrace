from datetime import date
import json
from pathlib import Path
import shutil
import sqlite3
from uuid import uuid4

import faiss
import numpy as np

from schemas.enums import ToolName
from schemas.tool_calls import ToolCall
from data_pipeline.documents.build_file_index import main as build_kb_main
from tools.document_search.interface import document_search
from tools.document_search.sample_data import load_sample_chunks
from tools.document_search.search import bm25_search, filter_chunks
from tools.document_search.vector_search import clear_vector_cache


def test_filter_chunks_by_company_and_type() -> None:
    chunks = filter_chunks(
        load_sample_chunks(),
        company_id="000001.SZ",
        document_types=["regulatory_inquiry"],
        start_date=date(2023, 1, 1),
        end_date=date(2023, 12, 31),
    )
    assert len(chunks) == 1
    assert chunks[0].document_id == "INQUIRY-2023"


def test_bm25_search_returns_relevant_chunk() -> None:
    hits = bm25_search(load_sample_chunks(), query="存货 跌价准备 问询", top_k=3)
    assert hits
    assert hits[0].chunk.company_id == "000001.SZ"
    assert "存货" in hits[0].chunk.text


def test_document_search_tool_returns_evidence_in_explicit_demo_mode(monkeypatch) -> None:
    monkeypatch.setenv("FINTRACE_KB_PATH", str(Path(".tmp_tests") / "missing-demo.sqlite"))
    monkeypatch.setenv("FINTRACE_DOCUMENT_SEARCH_DEMO_MODE", "true")
    result = document_search(
        ToolCall(
            tool_call_id="CALL-001",
            tool_name=ToolName.DOCUMENT_SEARCH,
            arguments={
                "company_ids": ["000001.SZ"],
                "query": "存货 跌价准备 问询",
                "document_types": ["annual_report_note", "audit_report", "regulatory_inquiry"],
                "top_k": 3,
            },
            reason="test",
        )
    )
    assert result.status.value == "success"
    assert result.data["hits"]
    assert result.evidence
    assert result.data["source"] == "sample"


def test_document_search_fails_when_kb_is_missing_outside_demo(monkeypatch) -> None:
    monkeypatch.setenv("FINTRACE_KB_PATH", str(Path(".tmp_tests") / "missing-production.sqlite"))
    monkeypatch.setenv("FINTRACE_DOCUMENT_SEARCH_DEMO_MODE", "false")
    result = document_search(_call({"query": "存货风险", "company_ids": ["000001.SZ"]}))
    assert result.status.value == "failed"
    assert result.error is not None
    assert result.error.error_type.value == "DATA_NOT_AVAILABLE"


def test_document_search_rejects_invalid_arguments(monkeypatch) -> None:
    monkeypatch.setenv("FINTRACE_KB_PATH", str(Path(".tmp_tests") / "missing-invalid.sqlite"))
    result = document_search(_call({"query": " ", "top_k": 999}))
    assert result.status.value == "failed"
    assert result.error is not None
    assert result.error.error_type.value == "INVALID_ARGUMENT"


def test_filtered_vector_search_scores_only_matching_company(monkeypatch) -> None:
    test_root = Path(".tmp_tests") / f"document_vector_filter_{uuid4().hex}"
    raw_dir = test_root / "raw_documents"
    kb_dir = test_root / "knowledge_base"
    raw_dir.mkdir(parents=True, exist_ok=True)
    (raw_dir / "000001.SZ_2023-05-12_research_report.txt").write_text(
        "甲公司存货风险说明。", encoding="utf-8"
    )
    (raw_dir / "600519.SH_2023-05-12_research_report.txt").write_text(
        "乙公司存货风险说明。", encoding="utf-8"
    )
    try:
        assert build_kb_main(["--raw-dir", str(raw_dir), "--kb-dir", str(kb_dir)]) == 0
        _write_test_vectors(kb_dir)
        monkeypatch.setenv("FINTRACE_KB_PATH", str(kb_dir / "fintrace_kb.sqlite"))
        monkeypatch.setenv("FINTRACE_DOCUMENT_SEARCH_DEMO_MODE", "false")
        monkeypatch.setattr(
            "tools.document_search.vector_search.build_embedding_client",
            lambda: _FakeEmbeddingClient(),
        )
        clear_vector_cache()
        result = document_search(
            _call(
                {
                    "query": "存货风险",
                    "company_ids": ["000001.SZ"],
                    "mode": "vector",
                    "top_k": 1,
                    "pool_k": 1,
                }
            )
        )
        assert result.status.value == "success"
        assert result.data["hits"][0]["chunk"]["company_id"] == "000001.SZ"
        assert result.data["retrieval_debug"]["vector_strategy"] == "filtered_exact"
        assert result.metrics.vector_search_time_ms >= 0
    finally:
        clear_vector_cache()
        shutil.rmtree(test_root, ignore_errors=True)


def test_document_search_prefers_local_knowledge_base(monkeypatch) -> None:
    kb_path = Path("data/indexes/document_search/fintrace_kb.sqlite")
    monkeypatch.setenv("FINTRACE_KB_PATH", str(kb_path))
    result = document_search(
        ToolCall(
            tool_call_id="CALL-001",
            tool_name=ToolName.DOCUMENT_SEARCH,
            arguments={
                "company_ids": ["000001.SZ"],
                "query": "知识库专用词",
                "top_k": 3,
            },
            reason="test",
        )
    )
    assert result.status.value == "success"


def test_build_kb_and_search_txt_document(monkeypatch) -> None:
    test_root = Path(".tmp_tests") / f"document_kb_{uuid4().hex}"
    raw_dir = test_root / "raw_documents/test_company/inquiry_letter"
    kb_dir = test_root / "knowledge_base"
    raw_dir.mkdir(parents=True, exist_ok=True)
    source_file = raw_dir / "000001.SZ_2023-05-12_inquiry_letter.txt"
    source_file.write_text("监管问询函要求公司说明存货跌价准备是否充分。", encoding="utf-8")
    try:
        exit_code = build_kb_main(["--raw-dir", str(test_root / "raw_documents/test_company"), "--kb-dir", str(kb_dir)])
        assert exit_code == 0
        monkeypatch.setenv("FINTRACE_KB_PATH", str(kb_dir / "fintrace_kb.sqlite"))
        result = document_search(
            ToolCall(
                tool_call_id="CALL-001",
                tool_name=ToolName.DOCUMENT_SEARCH,
                arguments={
                    "company_ids": ["000001.SZ"],
                    "query": "存货 跌价准备 问询函",
                    "document_types": ["inquiry_letter"],
                    "top_k": 3,
                },
                reason="test",
            )
        )
        assert result.data["source"] == "knowledge_base"
        assert result.data["hits"]
        assert result.data["retrieval_debug"]["returned_hit_count"] == len(result.data["hits"])
        assert result.data["hits"][0]["retrieval"]["matched_by"]
        assert result.evidence[0].source.source_path.endswith("000001.SZ_2023-05-12_inquiry_letter.txt")
        assert result.evidence[0].fact["retrieval"]["final_score"] > 0
        assert (kb_dir / "parse_report.json").exists()
    finally:
        shutil.rmtree(test_root, ignore_errors=True)


def test_section_aware_chunking_writes_parse_report(monkeypatch) -> None:
    test_root = Path(".tmp_tests") / f"document_section_kb_{uuid4().hex}"
    raw_dir = test_root / "raw_documents/test_company/inquiry_letter"
    kb_dir = test_root / "knowledge_base"
    raw_dir.mkdir(parents=True, exist_ok=True)
    source_file = raw_dir / "000001.SZ_2023-05-12_inquiry_letter.txt"
    source_file.write_text("问题一：关于存货跌价准备\n\n请公司说明存货跌价准备是否充分。", encoding="utf-8")
    try:
        exit_code = build_kb_main(["--raw-dir", str(test_root / "raw_documents/test_company"), "--kb-dir", str(kb_dir)])
        assert exit_code == 0
        report = json.loads((kb_dir / "parse_report.json").read_text(encoding="utf-8"))
        assert report["summary"]["document_count"] == 1
        assert report["documents"][0]["section_count"] == 1
        monkeypatch.setenv("FINTRACE_KB_PATH", str(kb_dir / "fintrace_kb.sqlite"))
        result = document_search(
            ToolCall(
                tool_call_id="CALL-001",
                tool_name=ToolName.DOCUMENT_SEARCH,
                arguments={"company_ids": ["000001.SZ"], "query": "存货 跌价准备", "document_types": ["inquiry_letter"]},
                reason="test",
            )
        )
        assert result.data["hits"][0]["chunk"]["section"] == "问题一：关于存货跌价准备"
    finally:
        shutil.rmtree(test_root, ignore_errors=True)


def test_docx_table_is_ingested_as_searchable_text(monkeypatch) -> None:
    import docx

    test_root = Path(".tmp_tests") / f"document_docx_table_kb_{uuid4().hex}"
    raw_dir = test_root / "raw_documents/test_company/annual_report"
    kb_dir = test_root / "knowledge_base"
    raw_dir.mkdir(parents=True, exist_ok=True)
    source_file = raw_dir / "000001.SZ_2023-04-30_annual_report.docx"
    document = docx.Document()
    document.add_paragraph("年度报告附注")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "期末余额"
    table.cell(0, 2).text = "跌价准备"
    table.cell(1, 0).text = "库存商品"
    table.cell(1, 1).text = "210"
    table.cell(1, 2).text = "15"
    document.save(str(source_file))
    try:
        exit_code = build_kb_main(["--raw-dir", str(test_root / "raw_documents/test_company"), "--kb-dir", str(kb_dir)])
        assert exit_code == 0
        report = json.loads((kb_dir / "parse_report.json").read_text(encoding="utf-8"))
        assert report["summary"]["table_count"] == 1
        monkeypatch.setenv("FINTRACE_KB_PATH", str(kb_dir / "fintrace_kb.sqlite"))
        result = document_search(
            ToolCall(
                tool_call_id="CALL-001",
                tool_name=ToolName.DOCUMENT_SEARCH,
                arguments={"company_ids": ["000001.SZ"], "query": "库存商品 跌价准备", "document_types": ["annual_report"]},
                reason="test",
            )
        )
        assert result.data["hits"]
        assert "库存商品" in result.data["hits"][0]["chunk"]["text"]
    finally:
        shutil.rmtree(test_root, ignore_errors=True)


def test_build_kb_skip_unchanged_reports_skipped_files() -> None:
    test_root = Path(".tmp_tests") / f"document_skip_kb_{uuid4().hex}"
    raw_dir = test_root / "raw_documents/test_company/inquiry_letter"
    kb_dir = test_root / "knowledge_base"
    raw_dir.mkdir(parents=True, exist_ok=True)
    source_file = raw_dir / "000001.SZ_2023-05-12_inquiry_letter.txt"
    source_file.write_text("监管问询函要求公司说明存货跌价准备是否充分。", encoding="utf-8")
    try:
        first_exit = build_kb_main(["--raw-dir", str(test_root / "raw_documents/test_company"), "--kb-dir", str(kb_dir)])
        second_exit = build_kb_main(
            [
                "--raw-dir",
                str(test_root / "raw_documents/test_company"),
                "--kb-dir",
                str(kb_dir),
                "--skip-unchanged",
            ]
        )
        assert first_exit == 0
        assert second_exit == 0
        report = json.loads((kb_dir / "parse_report.json").read_text(encoding="utf-8"))
        assert report["summary"]["skipped_count"] == 1
        assert report["documents"][0]["parse_status"] == "skipped_unchanged"
    finally:
        shutil.rmtree(test_root, ignore_errors=True)


class _FakeEmbeddingClient:
    model = "text-embedding-v4"
    dimension = 2
    api_mode = "compatible"

    def embed_query(self, text: str) -> np.ndarray:
        del text
        return np.asarray([1.0, 0.0], dtype="float32")


def _call(arguments: dict) -> ToolCall:
    return ToolCall(
        tool_call_id="CALL-001",
        tool_name=ToolName.DOCUMENT_SEARCH,
        arguments=arguments,
        reason="test",
    )


def _write_test_vectors(kb_dir: Path) -> None:
    with sqlite3.connect(kb_dir / "fintrace_kb.sqlite") as conn:
        rows = conn.execute("SELECT chunk_id, company_id FROM chunks ORDER BY chunk_id").fetchall()
    vector_ids = [row[0] for row in rows]
    embeddings = np.asarray(
        [[1.0, 0.0] if row[1] == "000001.SZ" else [0.0, 1.0] for row in rows],
        dtype="float32",
    )
    np.save(kb_dir / "embeddings.npy", embeddings)
    index = faiss.IndexFlatIP(2)
    index.add(embeddings)
    faiss.write_index(index, str(kb_dir / "vector.faiss"))
    (kb_dir / "vector_ids.json").write_text(
        json.dumps(vector_ids, ensure_ascii=False), encoding="utf-8"
    )
    (kb_dir / "manifest.json").write_text(
        json.dumps(
            {
                "embedding": {
                    "model": "text-embedding-v4",
                    "dimension": 2,
                    "api_mode": "compatible",
                }
            }
        ),
        encoding="utf-8",
    )
