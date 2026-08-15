"""Convert the five competition datasets to JSONL and fetch announcement text.

The script intentionally uses only the project dependencies and Python's standard
library. In particular, XLSX files are read in streaming mode without openpyxl so
the 600k-row shareholder workbook does not need to be loaded into memory.
"""

from __future__ import annotations

import argparse
import csv
import io
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
import threading
import time
import zipfile
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, Iterator
from urllib.parse import urlparse
from xml.etree import ElementTree as ET


SPREADSHEET_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
DATE_FIELDS = {
    "ann_dt",
    "s_holder_enddate",
    "report_period",
    "collect_dt",
    "actual_ann_dt",
    "write_date",
    "publish_date",
    "opdate",
}
TEXT_FIELDS = {
    "session_id",
    "object_id",
    "id",
    "report_id",
    "statement_type",
    "comp_type_code",
    "opmode",
    "s_holder_holdercategory",
    "s_holder_sharecategory",
    "s_holder_sequence",
    "n_info_fcode",
    "rating_id",
    "tmstamp",
}
TEXT_SUFFIXES = ("_id", "_code", "code", "_name", "_title", "_link")
NUMBER_RE = re.compile(r"^[+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?$")
CELL_COLUMN_RE = re.compile(r"[A-Z]+")
DATE_8_RE = re.compile(r"^(\d{4})(\d{2})(\d{2})$")
DATETIME_14_RE = re.compile(r"^(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})$")
BLOCK_TAGS = {
    "article",
    "blockquote",
    "br",
    "div",
    "footer",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "header",
    "li",
    "main",
    "p",
    "section",
    "table",
    "td",
    "th",
    "tr",
}
ATTACHMENT_RE = re.compile(r"\.(?:pdf|docx?|rtf)(?:\s|\(|$)", re.IGNORECASE)
OLE_DOCUMENT_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")


@dataclass
class DatasetResult:
    source: str
    output: str
    rows: int = 0
    columns: list[str] = field(default_factory=list)
    statuses: Counter = field(default_factory=Counter)
    errors: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "source": self.source,
            "output": self.output,
            "rows": self.rows,
            "column_count": len(self.columns),
            "columns": self.columns,
            "statuses": dict(self.statuses),
            "errors": self.errors[:20],
        }


class VisibleTextParser(HTMLParser):
    """Small dependency-free HTML-to-text parser for announcement pages."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.hidden_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg"}:
            self.hidden_depth += 1
        elif not self.hidden_depth and tag in BLOCK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg"}:
            self.hidden_depth = max(0, self.hidden_depth - 1)
        elif not self.hidden_depth and tag in BLOCK_TAGS:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.hidden_depth:
            self.parts.append(data)

    def text(self) -> str:
        lines = []
        for line in "".join(self.parts).replace("\r", "\n").split("\n"):
            cleaned = re.sub(r"\s+", " ", line).strip()
            if cleaned:
                lines.append(cleaned)
        return "\n".join(lines)


def column_index(cell_reference: str) -> int:
    match = CELL_COLUMN_RE.match(cell_reference)
    if not match:
        raise ValueError(f"Invalid XLSX cell reference: {cell_reference}")
    index = 0
    for character in match.group(0):
        index = index * 26 + ord(character) - ord("A") + 1
    return index - 1


def load_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    path = "xl/sharedStrings.xml"
    if path not in archive.namelist():
        return []
    strings: list[str] = []
    with archive.open(path) as handle:
        for _, element in ET.iterparse(handle, events=("end",)):
            if element.tag == f"{SPREADSHEET_NS}si":
                strings.append("".join(node.text or "" for node in element.iter(f"{SPREADSHEET_NS}t")))
                element.clear()
    return strings


def iter_xlsx_rows(path: Path) -> Iterator[list[object | None]]:
    """Yield rows from the first worksheet while preserving sparse columns."""

    with zipfile.ZipFile(path) as archive:
        shared_strings = load_shared_strings(archive)
        sheet_path = "xl/worksheets/sheet1.xml"
        if sheet_path not in archive.namelist():
            raise ValueError(f"No first worksheet found in {path}")
        with archive.open(sheet_path) as handle:
            for _, row in ET.iterparse(handle, events=("end",)):
                if row.tag != f"{SPREADSHEET_NS}row":
                    continue
                values: dict[int, object | None] = {}
                for cell in row.findall(f"{SPREADSHEET_NS}c"):
                    index = column_index(cell.attrib.get("r", ""))
                    cell_type = cell.attrib.get("t")
                    value_node = cell.find(f"{SPREADSHEET_NS}v")
                    raw_value = value_node.text if value_node is not None else None
                    if cell_type == "s" and raw_value is not None:
                        value: object | None = shared_strings[int(raw_value)]
                    elif cell_type == "inlineStr":
                        value = "".join(
                            node.text or "" for node in cell.iter(f"{SPREADSHEET_NS}t")
                        )
                    elif cell_type == "b" and raw_value is not None:
                        value = raw_value == "1"
                    else:
                        value = raw_value
                    values[index] = value
                width = max(values, default=-1) + 1
                yield [values.get(index) for index in range(width)]
                row.clear()


def iter_xlsx_records(path: Path) -> tuple[list[str], Iterator[dict[str, object | None]]]:
    rows = iter_xlsx_rows(path)
    try:
        headers = [str(value).strip() for value in next(rows)]
    except StopIteration as exc:
        raise ValueError(f"Empty workbook: {path}") from exc

    def records() -> Iterator[dict[str, object | None]]:
        for row in rows:
            yield {
                header: row[index] if index < len(row) else None
                for index, header in enumerate(headers)
            }

    return headers, records()


def iter_csv_records(path: Path) -> tuple[list[str], Iterator[dict[str, str | None]]]:
    handle = path.open("r", encoding="utf-8-sig", newline="")
    reader = csv.DictReader(handle)
    headers = list(reader.fieldnames or [])

    def records() -> Iterator[dict[str, str | None]]:
        try:
            yield from reader
        finally:
            handle.close()

    return headers, records()


def normalize_date(value: object) -> str | None:
    if value is None or str(value).strip() == "":
        return None
    text = str(value).strip()
    match = DATE_8_RE.match(text)
    if match:
        try:
            return datetime.strptime(text, "%Y%m%d").date().isoformat()
        except ValueError:
            return text
    match = DATETIME_14_RE.match(text)
    if match:
        try:
            return datetime.strptime(text, "%Y%m%d%H%M%S").isoformat()
        except ValueError:
            return text
    if re.fullmatch(r"\d{5}(?:\.0+)?", text):
        serial = int(float(text))
        if 20_000 <= serial <= 80_000:
            return (datetime(1899, 12, 30) + timedelta(days=serial)).date().isoformat()
    return text


def is_text_field(name: str) -> bool:
    lowered = name.lower()
    return (
        lowered in TEXT_FIELDS
        or lowered.endswith(TEXT_SUFFIXES)
        or any(part in lowered for part in ("name", "title", "memo", "abstract", "author", "link"))
    )


def normalize_value(name: str, value: object) -> object | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    text = str(value).strip()
    if not text:
        return None
    if name.lower() in DATE_FIELDS:
        return normalize_date(text)
    if is_text_field(name):
        return text
    if NUMBER_RE.fullmatch(text):
        if re.fullmatch(r"[+-]?\d+", text):
            if len(text.lstrip("+-")) > 1 and text.lstrip("+-").startswith("0"):
                return text
            return int(text)
        number = float(text)
        return number if math.isfinite(number) else None
    return text


def normalize_record(record: dict[str, object | None]) -> dict[str, object | None]:
    return {name: normalize_value(name, value) for name, value in record.items()}


def load_announcement_categories(path: Path) -> dict[str, str]:
    categories: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^(\d{10})\s+(.+?)\s*$", line)
        if match:
            categories[match.group(1)] = match.group(2)
    return categories


def announcement_record(
    raw_record: dict[str, object | None], categories: dict[str, str], data_dir: Path
) -> dict[str, object | None]:
    record = normalize_record(raw_record)
    category_codes = [code for code in str(record.get("n_info_fcode") or "").split("|") if code]
    html_url = str(record.get("n_info_windlink") or "") or None
    pdf_url = str(record.get("n_info_annlink") or "") or None
    document_id = str(record.get("id") or record.get("object_id") or "unknown").strip("{}")
    document_name = re.sub(r"[^A-Za-z0-9._-]+", "_", document_id) + ".txt"
    output_path = data_dir / "documents" / "announcements" / document_name
    try:
        document_path = output_path.relative_to(Path.cwd()).as_posix()
    except ValueError:
        document_path = output_path.as_posix()
    record.update(
        {
            "category_codes": category_codes,
            "category_names": [categories[code] for code in category_codes if code in categories],
            "source_urls": {"html": html_url, "pdf": pdf_url},
            "selected_url": None,
            "selected_format": None,
            "document_path": document_path,
            "download_status": "not_requested",
            "download_error": None,
        }
    )
    return record


def extract_html_text(content: bytes, encoding: str | None = None) -> str:
    charset = encoding or "utf-8"
    try:
        source = content.decode(charset)
    except (LookupError, UnicodeDecodeError):
        source = content.decode("utf-8", errors="replace")
    parser = VisibleTextParser()
    parser.feed(source)
    return parser.text()


def extract_pdf_text(content: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PDF parsing requires pymupdf from requirements.txt") from exc
    pages: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as document:
        for page in document:
            text = page.get_text("text").strip()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def extract_docx_text(content: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("DOCX parsing requires python-docx from requirements.txt") from exc
    document = docx.Document(io.BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def find_libreoffice() -> str | None:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if executable:
        return executable
    candidates = (
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    )
    return str(next((path for path in candidates if path.exists()), "")) or None


def extract_legacy_doc_text(content: bytes) -> str:
    executable = find_libreoffice()
    with tempfile.TemporaryDirectory(prefix="fintrace-doc-") as temporary_directory:
        directory = Path(temporary_directory)
        source = directory / "announcement.doc"
        output = directory / "announcement.txt"
        source.write_bytes(content)
        if executable:
            command = [
                    executable,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(directory),
                    str(source),
                ]
        elif os.name == "nt":
            powershell = shutil.which("powershell.exe") or shutil.which("powershell")
            converter = Path(__file__).with_name("convert_legacy_doc.ps1")
            if not powershell or not converter.is_file():
                raise RuntimeError("Legacy DOC parsing requires Microsoft Word or LibreOffice")
            command = [
                powershell,
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(converter),
                "-InputPath",
                str(source),
                "-OutputPath",
                str(output),
            ]
        else:
            raise RuntimeError("Legacy DOC parsing requires LibreOffice (soffice)")
        creation_flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        with _word_conversion_lock:
            process = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
                creationflags=creation_flags,
            )
        if process.returncode != 0 or not output.exists():
            message = process.stderr.strip() or process.stdout.strip() or "conversion failed"
            raise RuntimeError(f"Could not convert legacy DOC: {message}")
        return output.read_text(encoding="utf-8", errors="replace")


def detect_document_format(content: bytes, content_type: str, url: str) -> str:
    stripped = content.lstrip()
    lowered_type = content_type.lower()
    lowered_url = urlparse(url).path.lower()
    if content.startswith(b"%PDF"):
        return "pdf"
    if content.startswith(OLE_DOCUMENT_MAGIC):
        return "doc"
    if content.startswith(b"PK\x03\x04"):
        return "docx"
    if stripped.startswith((b"<", b"<!")) or "html" in lowered_type:
        return "html"
    if "application/pdf" in lowered_type:
        return "pdf"
    if lowered_url.endswith(".docx"):
        return "docx"
    if lowered_url.endswith(".doc"):
        return "doc"
    if lowered_url.endswith(".pdf"):
        return "pdf"
    return "html"


def is_html_metadata_only(text: str) -> bool:
    cleaned = text.strip()
    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    return len(cleaned) < 300 and len(lines) <= 8 and any(
        ATTACHMENT_RE.search(line) for line in lines
    )


_thread_local = threading.local()
_word_conversion_lock = threading.Lock()


def request_session(retries: int):
    import requests
    from requests.adapters import HTTPAdapter
    from urllib3.util.retry import Retry

    session = getattr(_thread_local, "session", None)
    session_retries = getattr(_thread_local, "session_retries", None)
    if session is None or session_retries != retries:
        session = requests.Session()
        retry = Retry(
            total=retries,
            connect=retries,
            read=retries,
            backoff_factor=1.0,
            status_forcelist=(429, 500, 502, 503, 504),
            allowed_methods=("GET",),
        )
        session.mount("http://", HTTPAdapter(max_retries=retry))
        session.mount("https://", HTTPAdapter(max_retries=retry))
        session.headers["User-Agent"] = "FinTrace-data-preprocessor/1.0"
        _thread_local.session = session
        _thread_local.session_retries = retries
    return session


def fetch_text(url: str, expected_format: str, timeout: float, retries: int = 2) -> tuple[str, str]:
    started = time.monotonic()
    with request_session(retries).get(
        url, timeout=(min(10.0, timeout), min(15.0, timeout)), stream=True
    ) as response:
        response.raise_for_status()
        chunks: list[bytes] = []
        size = 0
        for chunk in response.iter_content(chunk_size=64 * 1024):
            if not chunk:
                continue
            chunks.append(chunk)
            size += len(chunk)
            if size > 50 * 1024 * 1024:
                raise ValueError("Downloaded document exceeded 50 MiB limit")
            if time.monotonic() - started > timeout:
                raise TimeoutError(f"Download exceeded {timeout:g} second total timeout")
        content = b"".join(chunks)
        content_type = response.headers.get("Content-Type", "").lower()
        response_encoding = response.encoding
    actual_format = detect_document_format(content, content_type, url)
    if actual_format == "pdf":
        text = extract_pdf_text(content)
    elif actual_format == "docx":
        text = extract_docx_text(content)
    elif actual_format == "doc":
        text = extract_legacy_doc_text(content)
    else:
        text = extract_html_text(content, response_encoding)
        if is_html_metadata_only(text):
            raise ValueError("Downloaded HTML contained metadata and an attachment link but no body")
    if len(text.strip()) < 80:
        raise ValueError(f"Downloaded {actual_format} contained too little extractable text")
    return text.strip(), actual_format


def download_announcement(
    record: dict[str, object | None],
    data_dir: Path,
    timeout: float,
    overwrite: bool,
    attachment_only: bool = False,
    retries: int = 2,
) -> dict[str, object | None]:
    output_path = Path(str(record["document_path"]))
    if output_path.exists() and output_path.stat().st_size > 0 and not overwrite:
        record["download_status"] = "existing"
        return record

    source_urls = record["source_urls"]
    assert isinstance(source_urls, dict)
    candidates = [("attachment", source_urls.get("pdf"))] if attachment_only else [
        ("html", source_urls.get("html")),
        ("attachment", source_urls.get("pdf")),
    ]
    errors: list[str] = []
    for expected_format, url in candidates:
        if not url:
            continue
        try:
            text, actual_format = fetch_text(str(url), expected_format, timeout, retries=retries)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            temporary_path = output_path.with_suffix(".txt.tmp")
            temporary_path.write_text(text + "\n", encoding="utf-8")
            os.replace(temporary_path, output_path)
            record["selected_url"] = url
            record["selected_format"] = actual_format
            record["download_status"] = "success"
            record["download_error"] = None
            record["index_document"] = True
            return record
        except Exception as exc:  # Each failed source should fall back to the next one.
            errors.append(f"{expected_format}: {type(exc).__name__}: {exc}")
    record["download_status"] = "failed"
    record["download_error"] = " | ".join(errors) or "No usable source URL"
    record["index_document"] = False
    return record


def announcement_repair_reason(record: dict[str, object | None]) -> str | None:
    status = str(record.get("download_status") or "")
    if status in {"failed", "content_unavailable"}:
        return "download_failed"
    if record.get("selected_format") != "html":
        return None
    document_path = Path(str(record.get("document_path") or ""))
    if not document_path.is_file():
        return "document_missing"
    text = document_path.read_text(encoding="utf-8", errors="replace")
    if is_html_metadata_only(text):
        return "html_metadata_only"
    return None


def is_timeout_failure(record: dict[str, object | None]) -> bool:
    if record.get("download_status") != "content_unavailable":
        return False
    error = str(record.get("download_error") or "").lower()
    return "timeout" in error or "timed out" in error


def repair_announcement(
    record: dict[str, object | None],
    reason: str,
    data_dir: Path,
    timeout: float,
    retries: int,
) -> dict[str, object | None]:
    repaired = dict(record)
    if reason == "html_metadata_only":
        repaired["html_status"] = "metadata_only"
    result = download_announcement(
        repaired,
        data_dir=data_dir,
        timeout=timeout,
        overwrite=True,
        attachment_only=reason == "html_metadata_only",
        retries=retries,
    )
    if result["download_status"] == "success":
        if reason == "html_metadata_only":
            result["repair_status"] = "success_attachment_fallback"
            result["attachment_status"] = "success"
        elif result.get("selected_format") == "html":
            result["repair_status"] = "success_html_retry"
            result["html_status"] = "success"
        else:
            result["repair_status"] = "success_attachment_retry"
            result["attachment_status"] = "success"
        result["repaired_at"] = datetime.now().astimezone().isoformat()
    else:
        result["download_status"] = "content_unavailable"
        result["repair_status"] = "failed"
        result["repaired_at"] = datetime.now().astimezone().isoformat()
    return result


def write_jsonl_atomic(path: Path, records: Iterable[dict]) -> None:
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8", newline="\n") as handle:
        for record in records:
            handle.write(
                json.dumps(record, ensure_ascii=False, allow_nan=False, separators=(",", ":"))
                + "\n"
            )
    os.replace(temporary_path, path)


def load_review_issues(data_dir: Path) -> dict[str, str]:
    review_path = data_dir / "announcement_download_review.csv"
    if not review_path.is_file():
        return {}
    with review_path.open("r", encoding="utf-8-sig", newline="") as handle:
        return {
            str(row["id"]): str(row["issue"])
            for row in csv.DictReader(handle)
            if row.get("id") and row.get("issue")
        }


def salvage_existing_repair(record: dict, review_reason: str) -> bool:
    if review_reason != "html_metadata_only":
        return False
    document_path = Path(str(record.get("document_path") or ""))
    if not document_path.is_file():
        return False
    text = document_path.read_text(encoding="utf-8", errors="replace").strip()
    if len(text) < 80 or is_html_metadata_only(text):
        return False
    source_urls = record.get("source_urls") or {}
    record["selected_url"] = source_urls.get("pdf")
    suffix = Path(urlparse(str(source_urls.get("pdf") or "")).path).suffix.lower()
    record["selected_format"] = suffix.lstrip(".") or "attachment"
    record["download_status"] = "success"
    record["download_error"] = None
    record["index_document"] = True
    record["html_status"] = "metadata_only"
    record["attachment_status"] = "success"
    record["repair_status"] = "success_attachment_fallback"
    record["repaired_at"] = datetime.now().astimezone().isoformat()
    return True


def repair_announcements(
    data_dir: Path,
    workers: int,
    timeout: float,
    retries: int,
    timeout_only: bool = False,
) -> dict:
    announcements_path = data_dir / "jsonl" / "announcements.jsonl"
    if not announcements_path.is_file():
        raise FileNotFoundError(f"Announcement JSONL does not exist: {announcements_path}")

    review_issues = load_review_issues(data_dir)
    records: list[dict] = []
    targets: list[tuple[int, dict, str]] = []
    with announcements_path.open("r", encoding="utf-8") as handle:
        for index, line in enumerate(handle):
            record = json.loads(line)
            records.append(record)
            if timeout_only:
                if not is_timeout_failure(record):
                    continue
                reason = (
                    "html_metadata_only"
                    if record.get("html_status") == "metadata_only"
                    else "download_failed"
                )
                targets.append((index, record, reason))
                continue
            review_reason = review_issues.get(str(record.get("id")))
            if review_reason and salvage_existing_repair(record, review_reason):
                continue
            already_repaired = str(record.get("repair_status") or "").startswith("success_")
            reason = None if already_repaired else (review_reason or announcement_repair_reason(record))
            if reason:
                targets.append((index, record, reason))

    reasons = Counter(reason for _, _, reason in targets)
    write_jsonl_atomic(announcements_path, records)
    print(
        f"Repairing {len(targets):,} announcements "
        f"({', '.join(f'{name}={count}' for name, count in sorted(reasons.items()))})",
        flush=True,
    )

    def run_target(target: tuple[int, dict, str]) -> tuple[int, str, dict]:
        index, record, reason = target
        return index, reason, repair_announcement(
            record, reason, data_dir, timeout, retries=retries
        )

    results: list[tuple[int, str, dict]] = []
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = [executor.submit(run_target, target) for target in targets]
        for completed, future in enumerate(as_completed(futures), start=1):
            result = future.result()
            results.append(result)
            index, _, record = result
            records[index] = record
            if completed % 10 == 0:
                write_jsonl_atomic(announcements_path, records)
            if completed % 25 == 0 or completed == len(targets):
                print(f"  repaired {completed:,}/{len(targets):,}", flush=True)

    repair_statuses = Counter()
    errors: list[dict] = []
    for index, reason, record in results:
        repair_statuses[str(record.get("repair_status"))] += 1
        if record.get("repair_status") == "failed":
            errors.append(
                {
                    "id": record.get("id"),
                    "title": record.get("n_info_title"),
                    "reason": reason,
                    "error": record.get("download_error"),
                }
            )

    write_jsonl_atomic(announcements_path, records)

    remaining = Counter()
    for record in records:
        reason = announcement_repair_reason(record)
        if reason:
            remaining[reason] += 1
    remaining_timeouts = [record for record in records if is_timeout_failure(record)]
    report = {
        "finished_at": datetime.now().astimezone().isoformat(),
        "announcement_jsonl": str(announcements_path),
        "selected": len(targets),
        "selected_reasons": dict(reasons),
        "repair_statuses": dict(repair_statuses),
        "remaining_issues": dict(remaining),
        "timeout_only": timeout_only,
        "remaining_timeouts": len(remaining_timeouts),
        "errors": errors,
    }
    report_name = (
        "announcement_timeout_retry_report.json"
        if timeout_only
        else "announcement_repair_report.json"
    )
    report_path = data_dir / report_name
    temporary_report = report_path.with_suffix(".json.tmp")
    temporary_report.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    os.replace(temporary_report, report_path)
    if timeout_only:
        write_timeout_csv(data_dir / "announcement_timeout_remaining.csv", remaining_timeouts)
    return report


def write_timeout_csv(path: Path, records: Iterable[dict]) -> None:
    fieldnames = [
        "id",
        "stock_code",
        "announcement_date",
        "title",
        "html_url",
        "attachment_url",
        "document_path",
        "error",
    ]
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for record in records:
            source_urls = record.get("source_urls") or {}
            writer.writerow(
                {
                    "id": record.get("id"),
                    "stock_code": record.get("s_info_windcode"),
                    "announcement_date": record.get("ann_dt"),
                    "title": record.get("n_info_title"),
                    "html_url": source_urls.get("html"),
                    "attachment_url": source_urls.get("pdf"),
                    "document_path": record.get("document_path"),
                    "error": record.get("download_error"),
                }
            )
    os.replace(temporary_path, path)


def limited(records: Iterable[dict], limit: int | None) -> Iterator[dict]:
    for index, record in enumerate(records):
        if limit is not None and index >= limit:
            break
        yield record


def write_records(path: Path, records: Iterable[dict], result: DatasetResult) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8", newline="\n") as handle:
        for record in records:
            handle.write(json.dumps(record, ensure_ascii=False, allow_nan=False, separators=(",", ":")))
            handle.write("\n")
            result.rows += 1
            status = record.get("download_status")
            if status:
                result.statuses[str(status)] += 1
                if status == "failed" and len(result.errors) < 20:
                    result.errors.append(
                        f"{record.get('id') or record.get('object_id')}: {record.get('download_error')}"
                    )
    os.replace(temporary_path, path)


def convert_table(
    source: Path,
    output: Path,
    data_dir: Path,
    limit: int | None,
    categories: dict[str, str] | None = None,
    download_documents: bool = False,
    workers: int = 4,
    timeout: float = 20.0,
    overwrite_documents: bool = False,
) -> DatasetResult:
    if source.suffix.lower() == ".xlsx":
        headers, raw_records = iter_xlsx_records(source)
    else:
        headers, raw_records = iter_csv_records(source)
    result = DatasetResult(str(source), str(output), columns=headers)
    rows: Iterable[dict] = limited(raw_records, limit)

    if categories is not None:
        rows = (announcement_record(record, categories, data_dir) for record in rows)
        if download_documents:
            executor = ThreadPoolExecutor(max_workers=workers)
            rows = executor.map(
                lambda record: download_announcement(
                    record, data_dir, timeout=timeout, overwrite=overwrite_documents
                ),
                rows,
            )
            try:
                write_records(output, rows, result)
            finally:
                executor.shutdown(wait=True, cancel_futures=True)
            return result
    else:
        rows = (normalize_record(record) for record in rows)

    write_records(output, rows, result)
    return result


def find_source_root(explicit_path: Path | None) -> Path:
    if explicit_path:
        path = explicit_path.resolve()
        if not path.is_dir():
            raise FileNotFoundError(f"Input directory does not exist: {path}")
        return path
    matches = sorted(path for path in Path.cwd().glob("14-*") if path.is_dir())
    if len(matches) != 1:
        raise FileNotFoundError(
            f"Expected exactly one directory starting with '14-', found {len(matches)}"
        )
    return matches[0].resolve()


def find_one(directory: Path, pattern: str) -> Path:
    matches = sorted(path for path in directory.glob(pattern) if not path.name.startswith("~$"))
    if len(matches) != 1:
        raise FileNotFoundError(f"Expected one {pattern!r} in {directory}, found {len(matches)}")
    return matches[0]


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-dir", type=Path, help="Competition directory starting with 14-")
    parser.add_argument("--data-dir", type=Path, default=Path("data"), help="Output root")
    parser.add_argument(
        "--download-documents",
        action="store_true",
        help="Download announcement HTML first and fall back to PDF",
    )
    parser.add_argument(
        "--repair-downloads",
        action="store_true",
        help="Repair failed or metadata-only announcement documents in existing output",
    )
    parser.add_argument("--workers", type=int, default=4, help="Concurrent download workers")
    parser.add_argument("--timeout", type=float, default=45.0, help="HTTP timeout in seconds")
    parser.add_argument("--retries", type=int, default=2, help="HTTP retries per source URL")
    parser.add_argument(
        "--overwrite-documents", action="store_true", help="Download existing announcement text again"
    )
    parser.add_argument(
        "--limit", type=int, help="Process at most this many rows per output (for verification)"
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.workers < 1:
        raise ValueError("--workers must be at least 1")
    if args.retries < 0:
        raise ValueError("--retries cannot be negative")
    if args.limit is not None and args.limit < 1:
        raise ValueError("--limit must be at least 1")
    if args.repair_downloads and args.download_documents:
        raise ValueError("Use --repair-downloads separately from --download-documents")

    data_dir = args.data_dir.resolve()
    if args.repair_downloads:
        report = repair_announcements(
            data_dir, workers=args.workers, timeout=args.timeout, retries=args.retries
        )
        print(
            f"Repair complete: {report['repair_statuses']}; remaining: {report['remaining_issues']}",
            flush=True,
        )
        return 0

    source_root = find_source_root(args.input_dir)
    jsonl_dir = data_dir / "jsonl"
    categories = load_announcement_categories(find_one(source_root / "3", "*.txt"))
    jobs = [
        (find_one(source_root / "1", "*.xlsx"), jsonl_dir / "questions.jsonl", None),
        (find_one(source_root / "2", "*.xlsx"), jsonl_dir / "shareholders.jsonl", None),
        (
            find_one(source_root / "3", "*.xlsx"),
            jsonl_dir / "announcements.jsonl",
            categories,
        ),
        (
            find_one(source_root / "4", "asharebalancesheet_*.csv"),
            jsonl_dir / "balance_sheets.jsonl",
            None,
        ),
        (
            find_one(source_root / "4", "ashareincome_*.csv"),
            jsonl_dir / "income_statements.jsonl",
            None,
        ),
        (
            find_one(source_root / "4", "asharecashflow_*.csv"),
            jsonl_dir / "cashflows.jsonl",
            None,
        ),
        (
            find_one(source_root / "5", "rr_main_*.csv"),
            jsonl_dir / "research_reports.jsonl",
            None,
        ),
    ]

    started_at = datetime.now().astimezone()
    results: list[DatasetResult] = []
    for source, output, announcement_categories in jobs:
        print(f"Converting {source.name} -> {output}", flush=True)
        result = convert_table(
            source=source,
            output=output,
            data_dir=data_dir,
            limit=args.limit,
            categories=announcement_categories,
            download_documents=args.download_documents and announcement_categories is not None,
            workers=args.workers,
            timeout=args.timeout,
            overwrite_documents=args.overwrite_documents,
        )
        results.append(result)
        print(f"  wrote {result.rows:,} rows", flush=True)

    finished_at = datetime.now().astimezone()
    report = {
        "source_root": str(source_root),
        "data_root": str(data_dir),
        "started_at": started_at.isoformat(),
        "finished_at": finished_at.isoformat(),
        "download_documents": args.download_documents,
        "limited_to_rows": args.limit,
        "datasets": {Path(result.output).stem: result.as_dict() for result in results},
        "total_rows": sum(result.rows for result in results),
    }
    report_path = data_dir / "preprocess_report.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_report = report_path.with_suffix(".json.tmp")
    temporary_report.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    os.replace(temporary_report, report_path)
    print(f"Report written to {report_path}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
