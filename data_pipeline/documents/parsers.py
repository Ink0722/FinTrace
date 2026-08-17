from pathlib import Path


def parse_document(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return parse_text_file(path)
    if suffix == ".docx":
        return parse_docx_file(path)
    if suffix == ".pdf":
        return parse_pdf_file(path)
    raise ValueError(f"Unsupported document type: {path.suffix}")


def parse_text_file(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [{"page": None, "text": text}]


def parse_docx_file(path: Path) -> list[dict]:
    import docx

    document = docx.Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_texts = [format_table(table_index, table) for table_index, table in enumerate(document.tables, start=1)]
    text_parts = [*paragraphs, *table_texts]
    return [{"page": None, "text": "\n\n".join(part for part in text_parts if part), "table_count": len(table_texts)}]


def parse_pdf_file(path: Path) -> list[dict]:
    import fitz

    pages: list[dict] = []
    with fitz.open(str(path)) as document:
        for index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            table_texts = extract_pdf_table_texts(page)
            if table_texts:
                text = f"{text}\n\n" + "\n\n".join(table_texts)
            if text:
                pages.append({"page": index, "text": text, "table_count": len(table_texts)})
    return pages


def format_table(table_index: int, table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    lines = [f"表格 {table_index}："]
    if rows:
        lines.append("列：" + " | ".join(rows[0]))
    for row in rows[1:]:
        lines.append("行：" + " | ".join(row))
    return "\n".join(lines)


def extract_pdf_table_texts(page) -> list[str]:
    find_tables = getattr(page, "find_tables", None)
    if not find_tables:
        return []
    try:
        tables = find_tables()
    except Exception:
        return []
    texts: list[str] = []
    for table_index, table in enumerate(getattr(tables, "tables", []), start=1):
        try:
            rows = table.extract()
        except Exception:
            continue
        if not rows:
            continue
        lines = [f"表格 {table_index}：", "列：" + " | ".join(str(cell or "").strip() for cell in rows[0])]
        for row in rows[1:]:
            lines.append("行：" + " | ".join(str(cell or "").strip() for cell in row))
        texts.append("\n".join(lines))
    return texts
